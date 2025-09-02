# nlp/parser_resume.py

import os
import re
from typing import Dict, List, Optional

import numpy as np
import pdfplumber
import docx
from rapidfuzz import fuzz
from striprtf.striprtf import rtf_to_text
from torch import cosine_similarity

from nlp.vacancy_parcer import parse_vacancy

# Ленивая загрузка spaCy моделей
_nlp_cache = {}

def get_nlp(lang: str = "ru"):
    """Ленивая загрузка spaCy модели. Если модель не установлена — возвращает None."""
    import spacy
    if lang not in _nlp_cache:
        model_name = "ru_core_news_lg" if lang == "ru" else "en_core_web_sm"
        try:
            _nlp_cache[lang] = spacy.load(model_name)
        except Exception as e:
            print(f"[WARN] spaCy model '{model_name}' not found ({e}). Skills extraction will be limited.")
            _nlp_cache[lang] = None
    return _nlp_cache[lang]


def _normalize_text(text: str) -> str:
    """Приводим текст к нижнему регистру, убираем лишние пробелы и спецсимволы."""
    text = text.lower()
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\s+", " ", text)
    text = text.strip()
    return text


def extract_text_from_rtf(file_path: str) -> str:
    """Извлекает текст из RTF файла через striprtf."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        rtf_content = f.read()
    text = rtf_to_text(rtf_content)
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Извлекает текст из DOCX файла, включая таблицы."""
    doc = docx.Document(file_path)
    full_text = []

    # Текст из параграфов
    for p in doc.paragraphs:
        if p.text.strip():
            full_text.append(p.text)

    # Текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    return "\n".join(full_text)


def extract_text_from_pdf(file_path: str) -> str:
    """Извлекает текст из PDF файла через pdfplumber."""
    with pdfplumber.open(file_path) as pdf:
        pages = [page.extract_text() for page in pdf.pages]
        text = "\n".join(filter(None, pages))
    return text


def extract_text_from_file(file_path: str) -> str:
    """Извлекает текст из PDF, DOCX или RTF файла."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext == ".rtf":
        text = extract_text_from_rtf(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    return _normalize_text(text)





def dedupe_text_combined(
        text: str,
        min_words: int = 3,
        chunk_size: int = 7,
        fuzzy_threshold: int = 82,
        nlp_threshold: float = 0.90,
        lang: str = "ru",
        overlap: Optional[int] = 5
) -> str:
    """
    Профессиональная очистка текста от повторов:
    - min_words: минимальное количество слов в сегменте
    - chunk_size: количество слов в сегменте
    - fuzzy_threshold: % схожести для быстрого фильтра
    - nlp_threshold: косинусное сходство spaCy
    - lang: 'ru' или 'en'
    - overlap: количество слов для перекрытия при скользящем окне
    """
    nlp = get_nlp(lang)
    if not nlp:
        print("[WARN] NLP модель не доступна, будет использован только fuzzy фильтр.")

    # Нормализация текста
    text = text.replace("—", "-").replace("–", "-").replace("―", "-")
    text = re.sub(r"\s+", " ", text.lower()).strip()
    words = text.split()

    if overlap is None:
        overlap = chunk_size // 2

    # Разбиваем на сегменты со скользящим окном
    segments = []
    i = 0
    while i < len(words):
        chunk = words[i:i + chunk_size]
        if len(chunk) >= min_words:
            segments.append(" ".join(chunk))
        i += chunk_size - overlap

    # 1️⃣ Быстрый fuzzy фильтр
    deduped_segments = []
    for seg in segments:
        if any(fuzz.partial_ratio(seg, ds) >= fuzzy_threshold for ds in deduped_segments):
            continue
        deduped_segments.append(seg)

    # 2️⃣ Косинусное сходство через spaCy
    if nlp:
        final_segments = []
        docs = [nlp(seg) for seg in deduped_segments]
        for doc in docs:
            if any(doc.similarity(fdoc) >= nlp_threshold for fdoc in final_segments):
                continue
            final_segments.append(doc)
        return " ".join([d.text for d in final_segments])

    # Если NLP не доступна, возвращаем результат только с fuzzy
    return " ".join(deduped_segments)


def _is_valid_skill_token(token_text: str) -> bool:
    """
    Проверка, что токен может быть техническим навыком.
    Разрешаем буквы, цифры и ограниченный набор спецсимволов (+ # . - /).
    """
    token_text = token_text.lower()
    if token_text in {"c++", "c#", ".net"}:
        return True
    return bool(re.match(r"^[a-zа-я0-9+#./-]+$", token_text, flags=re.IGNORECASE))



def extract_skills_from_text(text: str, vacancy_data: Dict, fuzzy_threshold: int = 75) -> List[Dict]:
    """
    Находит навыки из vacancy_data['requirements'] в тексте резюме.
    Возвращает список словарей: {"skill": ..., "match_type": "phrase|token|fuzzy", "score": int}.
    """
    if not vacancy_data or "requirements" not in vacancy_data:
        return []

    # Нормализуем текст один раз
    text_norm = text.lower()

    # Загружаем NLP (лениво)
    nlp_ru = get_nlp("ru")
    nlp_en = get_nlp("en")

    # Токены резюме (леммы + тех. токены)
    token_set: set = set()
    if nlp_ru:
        for t in nlp_ru(text_norm):
            if not t.is_stop and _is_valid_skill_token(t.text):
                token_set.add((t.lemma_ or t.text).lower())
    if nlp_en:
        for t in nlp_en(text_norm):
            if not t.is_stop and _is_valid_skill_token(t.text):
                token_set.add((t.lemma_ or t.text).lower())

    results: List[Dict] = []
    required_skills = [s for s in vacancy_data.get("requirements", []) if isinstance(s, str) and s.strip()]

    for skill in required_skills:
        sk_norm = skill.lower()

        # 1) Фразовое совпадение
        if sk_norm in text_norm:
            results.append({"skill": skill, "match_type": "phrase", "score": 100})
            continue

        # 2) Токен-совпадение (считаем леммы навыка только если до сюда дошли)
        skill_token_list: List[str] = []
        if nlp_ru:
            for t in nlp_ru(sk_norm):
                if _is_valid_skill_token(t.text):
                    skill_token_list.append((t.lemma_ or t.text).lower())
        if nlp_en:
            for t in nlp_en(sk_norm):
                if _is_valid_skill_token(t.text):
                    skill_token_list.append((t.lemma_ or t.text).lower())

        if skill_token_list and any(tok in token_set for tok in skill_token_list):
            results.append({"skill": skill, "match_type": "token", "score": 90})
            continue

        # 3) Fuzzy-совпадение (опечатки/варианты)
        fscore = int(fuzz.partial_ratio(sk_norm, text_norm))
        if fscore >= fuzzy_threshold:
            results.append({"skill": skill, "match_type": "fuzzy", "score": fscore})

    return results





# ------------------- Основная функция -------------------
def parse_resume(file_path: str, raw_vacancy: Dict = None) -> Dict:
    """
    Парсит резюме, нормализует текст, извлекает навыки и структурированный опыт.
    raw_vacancy — это исходный словарь вакансии из базы данных.
    """
    text = extract_text_from_file(file_path)
    text = dedupe_text_combined(text)

    # Нормализуем и парсим данные вакансии (если передана)
    vacancy_data = parse_vacancy(raw_vacancy) if raw_vacancy else None

    # Извлекаем навыки (если есть данные вакансии)
    skills_detailed = extract_skills_from_text(text, vacancy_data) if vacancy_data else []
    skills = [s["skill"] for s in skills_detailed]

    parsed = {
        "raw_text": text,
        "skills": skills,
        "skills_detailed": skills_detailed,
        "experience": [],
        "education": [],
    }
    return parsed




